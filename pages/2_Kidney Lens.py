import streamlit as st
import time
from azure.cognitiveservices.vision.customvision.prediction import CustomVisionPredictionClient
from msrest.authentication import ApiKeyCredentials
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

# Replace with your endpoint and prediction key
ENDPOINT = st.secrets["kidney_endpoint"]
PREDICTION_KEY = st.secrets["kidney_PREDICTION_KEY"]
API_TOKEN = st.secrets["kidney_API_TOKEN"]
account_id = st.secrets["kidney_account_id"]

# Create a prediction client
credentials = ApiKeyCredentials(in_headers={"Prediction-key": PREDICTION_KEY})
predictor = CustomVisionPredictionClient(ENDPOINT, credentials)

st.set_page_config(page_title="LifeTune", page_icon="💊")

database_endpoint="https://hvbajoria101.kintone.com/k/v1/record.json?"
database_headers={'X-Cybozu-API-Token':f"{st.secrets['kintone_key']}", 'Content-Type': 'application/json'}

doctors = []

# Fetching doctors from database
for i in range(1,3):
    database_data = {
        'app':1,
        'id':i
    }

    database_response = requests.get(f"{database_endpoint}", headers=database_headers, json=database_data)
    
    doctor={}
    doctor["name"]=database_response.json()["record"]["Text"]["value"]
    doctor["specialization"]=database_response.json()["record"]["Text_0"]["value"]
    doctor["location"]=database_response.json()["record"]["Text_1"]["value"]
    doctor["available_days"]=database_response.json()["record"]["Text_3"]["value"]
    doctor["contact"]=database_response.json()["record"]["Text_2"]["value"]
    doctors.append(doctor)

def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a friendly medical assistant" },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/meta/llama-2-7b-chat-int8", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("LifeTune_Chat.docx")

def runner():

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")
    doc = docx.Document()
    doc.save("LifeTune_Chat.docx")


    if user_input:
        if user_input:
            bot_reply = bot_response(user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on end button to remove chat", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("LifeTune_Chat.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            )

cystcauses = """
It's not clear what causes simple kidney cysts. One theory suggests that kidney cysts develop when the surface layer of the kidney weakens and forms a pouch. The pouch then fills with fluid, detaches and develops into a cyst.
"""
cystsymptoms = """
Simple kidney cysts rarely cause trouble, but growing giants can stir. If a cyst balloons, expect lurking aches in your back or side, feverish whispers, or pangs in your upper abdomen. Don't wait for a chorus of complaints - consult a doctor if any solo symptom starts singing too loud.
"""
cysttreat = """
If symptomatic, treatment for a simple kidney cyst may involve piercing and draining with a scarring solution, using a thin needle to drain and shrink the cyst, or surgery for larger cysts. Surgery is uncommon for simple cysts but may be considered for complex cysts with potential cancerous changes.
"""

stonecauses = """Kidney stones often have no definite, single cause, although several factors may increase your risk. Kidney stones form when your urine contains more crystal-forming substances — such as calcium, oxalate and uric acid — than the fluid in your urine can dilute. At the same time, your urine may lack substances that prevent crystals from sticking together, creating an ideal environment for kidney stones to form."""

stonesymptoms = """Kidney stones are often asymptomatic until they move within the kidney or into the ureters. If lodged, they can cause severe pain, block urine flow, and lead to symptoms like abdominal and groin pain, painful urination, and more. Additional signs may include discolored urine, cloudy or foul-smelling urine, increased frequency, nausea, vomiting, and fever if infection is present. The pain's location and intensity may vary as the stone progresses through the urinary tract.
"""
stonetreat = """
Kidney stones may require treatments like shock wave lithotripsy (ESWL) or surgery, depending on size and complications. ESWL uses sound waves to break stones, lasting 45-60 minutes, with potential side effects. Surgery options include percutaneous nephrolithotomy for large stones and ureteroscopy for smaller ones. In cases of calcium phosphate stones due to hyperparathyroidism, surgery addresses gland issues or associated conditions to prevent stone formation.
"""

tumorcauses = """
It's not clear what causes most kidney cancers. 
Doctors know that kidney cancer begins when some kidney cells develop changes (mutations) in their DNA. A cell's DNA contains the instructions that tell a cell what to do. The changes tell the cells to grow and divide rapidly. The accumulating abnormal cells form a tumor that can extend beyond the kidney. Some cells can break off and spread (metastasize) to distant parts of the body."""

tumorsymptoms = """Early-stage kidney cancer often lacks noticeable symptoms. Over time, signs like blood in urine, persistent back/side pain, appetite loss, weight loss, fatigue, and fever may emerge.
"""
tumortreat = """For most kidney cancers, surgery is the primary treatment, aiming to remove the cancer while preserving kidney function. Surgical options include radical nephrectomy, removing the entire kidney, and partial nephrectomy, removing the tumor and a small margin of surrounding healthy tissue. Both procedures can be performed through open, laparoscopic, or robotic-assisted approaches.
"""
st.markdown(
    """
        <style>
            [data-testid="stSidebarNav"] {
                background-repeat: no-repeat;                
            }
            [data-testid="stSidebarNav"]::before {
                content: "LifeTune";
                margin-left: 20px;
                margin-top: 20px;

                font-size: 30px;
                text-align: center;
                position: relative;
            }
        </style>
        """,
    unsafe_allow_html=True,
)

def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "LifeTune: Kidney Lens"
  
# left_co, cent_co,last_co = st.columns(3)
# with cent_co:
#     st.image("images/logo.png", width=200)

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
    
st.text(
    "Upload an image of a close up of your CT scan and we will tell you what is the disease you are suffering from"
)
# read images.zip as a binary file and put it into the button
with open("Kidney.zip", "rb") as fp:
    btn = st.download_button(
        label="Download test images",
        data=fp,
        file_name="Kidney.zip",
        mime="application/zip",
    )
image = st.file_uploader(
    "Upload Image", type=["jpg", "jpeg", "png", "webp"], accept_multiple_files=False
)

if image is not None:
    disp = False
    
    with image:
        st.image(image, caption="Your CT Scan", width=350)
        image_data = image.read()
        results = predictor.classify_image("748d2598-c0ce-481d-8e75-6a0d65dd5fdc", "Iteration1", image_data)
    disp = True
    
    c = st.image("loader.gif")
    time.sleep(3)
    c.empty()
    # Process and display the results
    if results.predictions:
        st.subheader("Prediction Results:")
        name="unknown"
        predict=0
        for prediction in results.predictions:
            if prediction.probability > predict and prediction.probability > 0.5:
                predict = prediction.probability
                name = prediction.tag_name

    if name!="unknown":
        st.success(f"Detected Kidney {name} with high confidence", icon='📃')
        if name == "Cyst":
            st.write(
                """
                Kidney cysts are round pouches of fluid that form on or in the kidneys. Kidney cysts can occur with disorders that may impair kidney function. But more often, kidney cysts are a type called simple kidney cysts. Simple kidney cysts aren't cancer and rarely cause problems.
                """
            )
            st.image("images/cyst.png", caption="Kidney cyst", width=350)
            st.write("More Info")

            tab1, tab2, tab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with tab1:
                st.write(cystcauses)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cysts/symptoms-causes/syc-20374134)"
                )
            with tab2:
                st.write(cystsymptoms)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cysts/symptoms-causes/syc-20374134)"
                )
            with tab3:
                st.write(cysttreat)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cysts/diagnosis-treatment/drc-20374138#:~:text=If%20a%20simple%20kidney%20cyst,whether%20your%20kidney%20cyst%20changes.)"
                )
            
            
            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"Kidney {name}"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif (
            name == "Stone"
        ):
            st.write(
                """
                Kidney stones (also called renal calculi, nephrolithiasis or urolithiasis) are hard deposits made of minerals and salts that form inside your kidneys. Passing kidney stones can be quite painful, but the stones usually cause no permanent damage if they're recognized in a timely fashion.
                """
            )
            col1, col2=st.columns(2)
            col1.image("images/male_stone.png", caption="male stone", width=350)
            col2.image("images/female_stone.png", caption="female stone", width=350)
            st.write("Known Carried Diseases")
            btab1, btab2, btab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with btab1:
                st.write(stonecauses)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-stones/symptoms-causes/syc-20355755#:~:text=Kidney%20stones%20form%20in%20your,and%20blood%20in%20your%20urine.)"
                )
            with btab2:
                st.write(stonesymptoms)
                st.write(
                    "More Info can be found on the [Mayo clinic Website](https://www.mayoclinic.org/diseases-conditions/kidney-stones/symptoms-causes/syc-20355755#:~:text=Kidney%20stones%20form%20in%20your,and%20blood%20in%20your%20urine.)"
                )
            with btab3:
                st.write(stonetreat)
                st.write(
                    "More Info can be found on the [Mayo clinic Website](https://www.mayoclinic.org/diseases-conditions/kidney-stones/diagnosis-treatment/drc-20355759)"
                )
            
            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"Kidney {name}"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')

            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif name == "Tumor":
            st.write(
                """
                Kidney cancer is cancer that begins in the kidneys. Your kidneys are two bean-shaped organs, each about the size of your fist. They're located behind your abdominal organs, with one kidney on each side of your spine.

                In adults, renal cell carcinoma is the most common type of kidney cancer. Other less common types of kidney cancer can occur. Young children are more likely to develop a kind of kidney cancer called Wilms' tumor.
                """
            )
            st.image("images/kidneycancer.png", caption="Kidney Cancer", width=350)
            st.write("Known Carried Diseases")
            ctab1, ctab2, ctab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with ctab1:
                st.write(tumorcauses)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cancer/symptoms-causes/syc-20352664)"
                )
            with ctab2:
                st.write(tumorsymptoms)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cancer/symptoms-causes/syc-20352664)"
                )
            with ctab3:
                st.write(tumortreat)
                st.write(
                    "More Info can be found on the [MAYO clinic website](https://www.mayoclinic.org/diseases-conditions/kidney-cancer/symptoms-causes/syc-20352664)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"Kidney {name}"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')

            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

    else:
        st.succcess("Feel Safe! No disease detected")
        st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
        first_run = st.session_state.get("first_run", True)

        if first_run:
            if st.button("Chat with AI Bot"):
                st.session_state.first_run = False
                runner()
        else:
            runner()


    
