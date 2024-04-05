import streamlit as st
import time
from azure.cognitiveservices.vision.customvision.prediction import CustomVisionPredictionClient
from msrest.authentication import ApiKeyCredentials
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

# Replace with your endpoint and prediction key
ENDPOINT = st.secrets["brain_ENDPOINT"]
PREDICTION_KEY = st.secrets["brain_PREDICTION_KEY"]
API_TOKEN = st.secrets["brain_API_TOKEN"]
account_id = st.secrets["brain_account_id"]

# Create a prediction client
credentials = ApiKeyCredentials(in_headers={"Prediction-key": PREDICTION_KEY})
predictor = CustomVisionPredictionClient(ENDPOINT, credentials)

st.set_page_config(page_title="LifeTune", page_icon="💊")
database_endpoint="https://hvbajoria101.kintone.com/k/v1/record.json?"
database_headers={'X-Cybozu-API-Token':f"{st.secrets['kintone_key']}", 'Content-Type': 'application/json'}

doctors = []

# Fetching doctors from database
for i in range(3,6):
    database_data = {
        'app':1,
        'id':i
    }

    database_response = requests.get(f"{database_endpoint}", headers=database_headers, json=database_data)
    
    doctor={}
    doctor["name"]=database_response.json()["record"]["Text"]["value"]
    doctor["specialization"]=database_response.json()["record"]["Text_0"]["value"]
    doctor["location"]=database_response.json()["record"]["Text_1"]["value"]
    doctor["available_days"]=database_response.json()["record"]["Text_3"]["value"]
    doctor["contact"]=database_response.json()["record"]["Text_2"]["value"]
    doctors.append(doctor)

def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a friendly medical assistant" },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/meta/llama-2-7b-chat-int8", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("LifeTune_Chat.docx")

def runner():

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")
    doc = docx.Document()
    doc.save("LifeTune_Chat.docx")


    if user_input:
        if user_input:
            bot_reply = bot_response(user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on end button to remove chat", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("LifeTune_Chat.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            )



gcauses = """
While the exact "why" of glioma remains hidden, research shines some light. Mutations in genes controlling growth and repair are likely culprits, sometimes inherited, sometimes arising spontaneously. Exposure to high-dose radiation like from atomic bombs also increases risk. These changes can cause normal cells to transform into malignant gliomas, either directly or by accumulating further mutations over time. These tumors lose their specialized functions and become genetically diverse, with even epigenetic tweaks influencing their growth. Understanding these complexities fuels the search for better diagnosis and treatment.
"""
gsymptoms = """
Though diverse in grade and location, gliomas share common symptoms like seizures, weakness, speech difficulties, and headaches. These arise from the tumor's pressure on brain tissue. Fatigue, vision changes, and psychological shifts can also occur. Malignant types may risk blood clots. Age at onset and survival vary heavily across glioma subtypes. Gene mutations in tumors help predict prognosis and guide treatment. Gliomas may increase in grade over time, with higher grades linked to lower survival. Their altered metabolism, vascularization, and invasive tendencies pose significant challenges.
"""
gtreat = """
Gliomas, brain tumors demanding meticulous care, necessitate comprehensive, multidisciplinary management. From emergency room presentation to surgical resection by skilled neurosurgeons, the journey involves radiologists, neuropathologists, and radiation/medical oncologists like neuro-oncologists, tailoring therapy based on tumor type, size, and patient specifics. Inoperable tumors may require subsequent chemotherapy with agents like temozolomide, while well-defined ones may be managed solely by surgery. Constant MRIs monitor progress, and additional specialists, such as neurologists and rehabilitation teams, may intervene as needed. Anti-seizure, anti-coagulant, and corticosteroid medications play crucial roles, with palliative care offering comfort when necessary. This intricate web of expertise ensures optimal care for each patient at every stage.
"""

mcauses = """Meningioma, a tumor arising from brain and spinal cord membranes, remains shrouded in mystery. While the exact cause escapes us, genetic mutations, particularly in the NF2 gene, likely play a role. These mutations disrupt essential tumor-suppressing proteins, allowing uncontrolled growth. Additionally, factors like survivin and vascularization boost proliferation, and aggressive forms might even lengthen their lifespan through telomere manipulation. Although the puzzle persists, researchers actively piece it together, paving the way for future understanding and treatment.
"""
msymptoms = """Meningiomas, primarily affecting older adults (median age 65), present varying symptoms based on tumor location and compromised brain/spine regions. Parasagittal tumors may cause headaches, seizures, limb weakness, vision issues, personality changes, or apathy. Cerebellar involvement can lead to gait imbalance, incoordination, eye movements, or hearing loss. Spinal meningiomas often cause motor or sensory issues in the legs, and bladder/bowel control problems. While benign, survival rates are high (80%+ at 5 years), with lower grades offering better prognoses and complete removals typically preventing recurrence. Incomplete resections or higher grades increase recurrence risk, especially within 15 years. Metastases, though rare, occur more frequently with grade 3 tumors.
"""
mtreat = """
Diagnosing meningioma requires a triplepronged approach: history & exam, medical imaging, and cellular analysis. Suspicious symptoms like seizures or neurological deficits trigger brain/spine imaging, with MRI preferred for initial tumor evaluation. While characteristic location hints at meningioma on scans, definitive diagnosis comes from analyzing tumor cells. This can involve a biopsy, but meningiomas are often surgically removed, achieving both diagnosis and treatment. Pathologists then analyze the removed cells under a microscope, often aided by immunohistochemistry for clearer identification. This comprehensive approach ensures accurate diagnosis and informs optimal treatment.
"""

pcauses = """
The exact causes of pituitary tumors, also known as pituitary adenomas, are not fully understood. However, certain factors may increase the risk of their development. These include genetic conditions like multiple endocrine neoplasia type 1 (MEN1) and Carney complex, as well as rare hereditary syndromes such as familial isolated pituitary adenoma. Hormonal imbalances, exposure to certain chemicals, and head injuries have also been suggested as potential contributing factors. However, in many cases, the underlying cause of pituitary tumors remains unknown."""

psymptoms = """Pituitary tumors can have diverse effects depending on their size, location, and hormone production. They can disrupt the normal functioning of the pituitary gland, leading to hormonal imbalances and associated symptoms. The specific effects can vary widely, ranging from vision problems and headaches due to pressure on nearby structures, to hormonal disturbances resulting in issues such as infertility, growth abnormalities, changes in body composition, and metabolic problems. The effects of pituitary tumors are highly dependent on the specific hormones involved and the individual's overall health.
"""
ptreat = """The treatment of pituitary tumors depends on several factors, including the tumor's size, hormone production, and the individual's overall health. Treatment options may include medication to regulate hormone levels, surgery to remove the tumor, radiation therapy to destroy tumor cells, or a combination of these approaches. The choice of treatment is determined by a multidisciplinary team of medical professionals and is tailored to the individual patient's needs and circumstances. Regular monitoring and follow-up care are often necessary to manage hormone levels, monitor tumor growth, and ensure optimal treatment outcomes.
"""
st.markdown(
    """
        <style>
            [data-testid="stSidebarNav"] {
                background-repeat: no-repeat;                
            }
            [data-testid="stSidebarNav"]::before {
                content: "LifeTune";
                margin-left: 20px;
                margin-top: 20px;

                font-size: 30px;
                text-align: center;
                position: relative;
            }
        </style>
        """,
    unsafe_allow_html=True,
)

def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "LifeTune: Brain Lens"
  
# left_co, cent_co,last_co = st.columns(3)
# with cent_co:
#     st.image("images/logo.png", width=200)

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
st.text(
    "Upload an image of a close up of a tumerous MRI scan and we will tell you what type it is."
)
# read images.zip as a binary file and put it into the button
with open("brain.zip", "rb") as fp:
    btn = st.download_button(
        label="Download test images",
        data=fp,
        file_name="brain.zip",
        mime="application/zip",
    )
image = st.file_uploader(
    "Upload Image", type=["jpg", "jpeg", "png", "webp"], accept_multiple_files=False
)

if image is not None:
    disp = False
    
    with image:
        st.image(image, caption="Your MRI Scan", width=350)
        image_data = image.read()
        results = predictor.classify_image("414083bc-72e5-4997-8145-ef52688393f7", "Iteration2", image_data)
    disp = True
    
    c = st.image("loader.gif")
    time.sleep(3)
    c.empty()
    # Process and display the results
    if results.predictions:
        st.subheader("Prediction Results:")
        name="unknown"
        predict=0
        for prediction in results.predictions:
            if prediction.probability > predict and prediction.probability > 0.5:
                predict = prediction.probability
                name = prediction.tag_name

    if name!="unknown":
        st.success(f"Detected {name} with high confidence", icon='📃')
        if name == "Glioma":
            st.write(
                """
                A glioma is a tumor of the central nervous system that arises from glial stem or progenitor cells. Glial cells are a type of cell widely present in the nervous system. Gliomas mostly occur in the brain and, rarely, in the spinal cord.
                """
            )
            st.image("images/glioma.webp", caption="Glioma", width=350)
            st.write("More Info")

            tab1, tab2, tab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with tab1:
                st.write(gcauses)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with tab2:
                st.write(gsymptoms)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with tab3:
                st.write(gtreat)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            
            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Brain Tumor"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif (
            name == "Meningioma"
        ):
            st.write(
                """
                A meningioma is a tumor arising from the meninges, which are membranes covering the brain and spinal cord. Excluding brain metastases from other organs, they are the most common intracranial tumor and are the most common primary brain tumor. 
                """
            )
            st.image("images/Meningioma.jfif", caption="Meningioma", width=350)
            st.write("Known Carried Diseases")
            btab1, btab2, btab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with btab1:
                st.write(mcauses)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with btab2:
                st.write(msymptoms)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with btab3:
                st.write(mtreat)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            
            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Brain Tumor"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif name == "Pituitary":
            st.write(
                """
                A rare pituitary tumor characterized by the presence of a pituitary adenoma that has metastasized either within the central nervous system, or to distant sites. The vast majority of pituitary carcinomas are hormonally active, most frequently with ACTH or prolactin production. 
                """
            )
            st.image("images/petu.jfif", caption="Pituitary", width=350)
            st.write("Known Carried Diseases")
            ctab1, ctab2, ctab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with ctab1:
                st.write(pcauses)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with ctab2:
                st.write(psymptoms)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
            with ctab3:
                st.write(ptreat)
                st.write(
                    "More Info can be found on the [National Organization for Rare Disorders (NORD)](https://rarediseases.org/)"
                )
           
            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Brain Tumor"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

    else:
        st.success("Feel safe! No tumor detected")
        st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
        first_run = st.session_state.get("first_run", True)

        if first_run:
            if st.button("Chat with AI Bot"):
                st.session_state.first_run = False
                runner()
        else:
            runner()
    
