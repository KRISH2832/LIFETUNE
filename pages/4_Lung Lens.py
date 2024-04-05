import streamlit as st
import time
from azure.cognitiveservices.vision.customvision.prediction import CustomVisionPredictionClient
from msrest.authentication import ApiKeyCredentials
import requests
import docx
from streamlit_extras.switch_page_button import switch_page

database_endpoint="https://hvbajoria101.kintone.com/k/v1/record.json?"
database_headers={'X-Cybozu-API-Token':f"{st.secrets['kintone_key']}", 'Content-Type': 'application/json'}

doctors = []

# Fetching doctors from database
for i in range(3,8):
    database_data = {
        'app':1,
        'id':i
    }

    database_response = requests.get(f"{database_endpoint}", headers=database_headers, json=database_data)
    
    doctor={}
    doctor["name"]=database_response.json()["record"]["Text"]["value"]
    doctor["specialization"]=database_response.json()["record"]["Text_0"]["value"]
    doctor["location"]=database_response.json()["record"]["Text_1"]["value"]
    doctor["available_days"]=database_response.json()["record"]["Text_3"]["value"]
    doctor["contact"]=database_response.json()["record"]["Text_2"]["value"]
    doctors.append(doctor)

# Replace with your endpoint and prediction key
ENDPOINT = st.secrets["lung_ENDPOINT"]
PREDICTION_KEY = st.secrets["lung_PREDICTION_KEY"]
API_TOKEN = st.secrets["lung_API_TOKEN"]
account_id = st.secrets["lung_account_id"]

# Create a prediction client
credentials = ApiKeyCredentials(in_headers={"Prediction-key": PREDICTION_KEY})
predictor = CustomVisionPredictionClient(ENDPOINT, credentials)

st.set_page_config(page_title="LifeTune", page_icon="💊")

def bot_response(question):
    API_BASE_URL = f"https://api.cloudflare.com/client/v4/accounts/{account_id}/ai/run/"
    headers = {"Authorization": f"Bearer {API_TOKEN}",'Content-Type': 'application/json'}

    def run(model, prompt):
        input = {
            "messages": [
            { "role": "system", "content": "You are a friendly medical assistant" },
            { "role": "user", "content": prompt }
            ]
        }
        response = requests.post(f"{API_BASE_URL}{model}", headers=headers, json=input)
        return response.json()

    output = run("@cf/meta/llama-2-7b-chat-int8", question)
    print(output)
    return output["result"]["response"]

def save_to_doc(conversation):
    doc = docx.Document()
    doc.add_heading("Chatbot Conversation", level=1)

    for user, bot in conversation:
        p_user = doc.add_paragraph()
        p_user.add_run("User: ").bold = True
        p_user.add_run(user)

        p_bot = doc.add_paragraph()
        p_bot.add_run("Bot: ").bold = True
        p_bot.add_run(bot)

    doc.save("LifeTune_Chat.docx")

def runner():

    if "conversation" not in st.session_state:
        st.session_state.conversation = []

    user_input = st.chat_input(placeholder="Your message")
    doc = docx.Document()
    doc.save("LifeTune_Chat.docx")

    if user_input:
        if user_input:
            bot_reply = bot_response(user_input)
            st.session_state.conversation.append((user_input, bot_reply))
        else:
            st.warning("Please enter a question.")
        
        save_to_doc(st.session_state.conversation)

    st.text("Conversation History:")
    for user, bot in st.session_state.conversation:
        text = st.chat_message("User")
        message = st.chat_message("Assistant")
        text.write(user)
        message.write(bot)
    
    if st.button("End Conversation"):
        st.session_state.conversation=[]
        st.session_state.knowledge=""
        st.session_state.first_run = True
        st.snow()
        st.success("Click on end button to remove chat", icon='✅')
        return
    
    st.download_button(
                label="Download Conversation",
                data=open("LifeTune_Chat.docx", "rb").read(),
                file_name="Conversation.docx",
                mime="application/octet-stream",
                help="Click to download the conversation."
            )

gcauses = """
There are a few factors that contribute to the development of this disease. The most common adenocarcinoma causes include:

Smoking- Tobacco use is the primary cause of adenocarcinoma and other types of cancer.

Toxin exposure- Harmful toxins in your home or work environment can also cause adenocarcinoma.

Previous radiation therapy- If you've had radiation therapy in the past, you have a higher risk of developing adenocarcinoma.
"""
gsymptoms = """
The first symptom is usually a chronic cough. You may cough up saliva and mucous with small amounts of blood. Other symptoms may include:

Difficulty breathing

Chest pain

Wheezing

Hoarseness

Loss of appetite

Weight loss
"""
gtreat = """
The treatment recommended for you will depend on the location, size and type of tumor. It also depends on whether or not the cancer has spread to other parts of your body. There are three main treatments for adenocarcinoma:

Surgery- Usually the first line of treatment for adenocarcinoma, surgery is done to remove cancer and some of the surrounding tissue.

Chemotherapy- This treatment involves using drugs to kill cancer cells. Chemotherapy may be used in a specific area or throughout your entire body.

Radiation therapy- Often used in combination with chemotherapy or surgery, radiation therapy uses imaging to target adenocarcinoma tumors and leave healthy tissues intact.
"""

mcauses = """Some of the causes and risk factors for squamous cell lung carcinoma include:

Smoking

Of all the causes of lung cancer, smoking is the most important. According to the National Cancer Institute, smokers are 10 times more likely to get any lung cancer than people who have smoked fewer than 100 cigarettes.

Radon exposure
The Environmental Protection Agency (EPA) lists radon as the second leading cause of lung cancer. It’s also the most common cause of lung cancer in nonsmokers.
Radon is a radioactive, odorless, invisible gas from rocks and soil. It’s a problem only in enclosed places, such as a house, because radon concentration is higher. People who smoke and are exposed to radon have a much higher risk for lung cancer.

Secondhand smoke exposure

Being exposed to secondhand smoke can increase your risk of lung cancer.


"""
msymptoms = """Many people don't experience squamous cell lung carcinoma symptoms until the cancer spreads. Some of the most common symptoms are:

persistent cough

bloody sputum

shortness of breath or wheezing

hoarseness

chest pain, especially when taking a deep breath or coughing

unexplained weight loss

decreased appetite

fatigue

It's also possible that people with squamous cell lung carcinoma may experience recurring lung infections. These can include pneumonia or bronchitis.


"""
mtreat = """
Treatment for squamous cell lung carcinoma depends on how advanced the cancer is, your ability to tolerate the side effects, and your overall health. Age isn’t usually a consideration.

The treatment you receive will be specific to your situation, but there are some general guidelines for the treatment of each stage.
"""

pcauses = """The exact causes of large cell carcinoma, like other types of lung cancer, are not fully understood, but several risk factors have been identified. Here are some of the key factors associated with the development of large cell carcinoma:
Smoking

Secondhand smoke

Genetic factors

Lung disease
"""
psymptoms = """ Large cell carcinoma and other forms of non-small cell lung cancer share the same symptoms, which may include:

persistent cough that gets worse over time

coughing up blood

trouble breathing

chest pain

wheezing

hoarseness

weight loss without trying

poor appetite

fatigue or tiredness

difficulty swallowing

swelling around the face or neck
"""
ptreat = """Here are some of the main types of treatment for non-small cell lung cancers like Large cell carcinoma.

Surgery

Surgery may be one of the first treatment steps if cancer was detected early and hasn't spread.

Radiation therapy

When tumors have grown too large to remove, or when trying to avoid removing large sections of lung, radiation therapy may be used to try and control cancer growth.

Chemotherapy

Chemotherapy involves the use of a number of medications that kill fast-growing cells like cancer. It can affect other cells that grow fast, too, like skin and hair.

Targeted therapy

With some cancers, like those caused by genetic mutations, special medications that target cancer cells but spare healthy cells may be used. These medications typically don't cure cancer but control growth and spread.

Immunotherapy

Immunotherapy is a growing area of medicine, where the immune system is programmed to fight a variety of conditions.



"""
st.markdown(
    """
        <style>
            [data-testid="stSidebarNav"] {
                background-repeat: no-repeat;                
            }
            [data-testid="stSidebarNav"]::before {
                content: "LifeTune";
                margin-left: 20px;
                margin-top: 20px;

                font-size: 30px;
                text-align: center;
                position: relative;
            }
        </style>
        """,
    unsafe_allow_html=True,
)

def gradient_text(text, color1, color2):
        gradient_css = f"""
        background: -webkit-linear-gradient(left, {color1}, {color2});
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: bold;
        font-size: 42px;
        """
        return f'<span style="{gradient_css}">{text}</span>'

color1 = "#0d3270"
color2 = "#0fab7b"
text = "LifeTune: Lung Lens"
  
# left_co, cent_co,last_co = st.columns(3)
# with cent_co:
#     st.image("images/logo.png", width=200)

styled_text = gradient_text(text, color1, color2)
st.write(f"<div style='text-align: center;'>{styled_text}</div>", unsafe_allow_html=True)
st.text(
    "Upload an image of a close up of a CT scan and we will tell you what type of Lung Cancer it is."
)
# read images.zip as a binary file and put it into the button
with open("lung.zip", "rb") as fp:
    btn = st.download_button(
        label="Download test images",
        data=fp,
        file_name="lung.zip",
        mime="application/zip",
    )
image = st.file_uploader(
    "Upload Image", type=["jpg", "jpeg", "png", "webp"], accept_multiple_files=False
)

if image is not None:
    disp = False
    
    with image:
        st.image(image, caption="Your CT Scan", width=350)
        image_data = image.read()
        results = predictor.classify_image("e9f1afe5-1690-4ce7-88a4-23a05ee90732", "Iteration1", image_data)
    disp = True
    
    c = st.image("loader.gif")
    time.sleep(3)
    c.empty()
    # Process and display the results
    if results.predictions:
        st.subheader("Prediction Results:")
        name="unknown"
        predict=0
        for prediction in results.predictions:
            if prediction.probability > predict and prediction.probability > 0.5:
                predict = prediction.probability
                name = prediction.tag_name

    if name!="unknown":
        st.success(f"Detected {name} with high confidence", icon='📃')
        if name == "adenocarcinoma":
            st.write(
                """
               Lung adenocarcinoma is a type of non-small cell lung cancer (NSCLC) that originates in the cells lining the small air sacs (alveoli) in the lungs. 
                """
            )
            st.image("images/Adenocarcinoma.png", caption="Adenocarcinoma", width=350)
            st.write("More Info")

            tab1, tab2, tab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with tab1:
                st.write(gcauses)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://my.clevelandclinic.org/health/diseases/21652-adenocarcinoma-cancers)"
                )
            with tab2:
                st.write(gsymptoms)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://my.clevelandclinic.org/health/diseases/21652-adenocarcinoma-cancers)"
                )
            with tab3:
                st.write(gtreat)
                st.write(
                    "More Info can be found on the [Mayo clinic website](https://my.clevelandclinic.org/health/diseases/21652-adenocarcinoma-cancers)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')

            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif (
            name == "squamous cell carcinoma"
        ):
            st.write(
                """
                Squamous cell carcinoma of the lungs, also known as squamous cell lung cancer, is a type of non-small cell lung cancer (NSCLC) that originates from the squamous epithelial cells lining the respiratory tract, particularly the bronchi. 
                """
            )
            st.image("images/Squamous cell carcinoma.png", caption="Squamous cell carcinoma", width=350)
            st.write("Known Carried Diseases")
            btab1, btab2, btab3 = st.tabs(
                ["Causes", "symptoms", "Treatment"]
            )
            with btab1:
                st.write(mcauses)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.healthline.com/health/lung-cancer/squamous-cell-lung-carcinoma#treatment)"
                )
            with btab2:
                st.write(msymptoms)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.healthline.com/health/lung-cancer/squamous-cell-lung-carcinoma#treatment)"
                )
            with btab3:
                st.write(mtreat)
                st.write(
                    "More Info can be found on the [Cancer Website](https://www.healthline.com/health/lung-cancer/squamous-cell-lung-carcinoma#treatment)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

        elif name == "large cell carcinoma":
            st.write(
                """
                Large cell carcinoma of the lungs is another type of non-small cell lung cancer (NSCLC).It is so named because the cancer cells are large and undifferentiated when examined under a microscope.
                """
            )
            st.image("images/Large Cell Carcinoma.png", caption="Large Cell Carcinoma", width=350)
            st.write("Known Carried Diseases")
            ctab1, ctab2, ctab3 = st.tabs(
                ["Causes", "Symptoms", "Treatment"]
            )
            with ctab1:
                st.write(pcauses)
                st.write(
                    "More Info can be found on the [Cancer.org](https://www.cancer.org/)"
                )
            with ctab2:
                st.write(psymptoms)
                st.write(
                    "More Info can be found on the [ Healthline.com](https://www.healthline.com/health/lung-cancer/large-cell-carcinoma#treatment)"
                )
            with ctab3:
                st.write(ptreat)
                st.write(
                    "More Info can be found on the [Healthline.com](https://www.healthline.com/health/lung-cancer/large-cell-carcinoma#treatment)"
                )

            book=st.button("Book Appointment with Doctor")
            if book:
                st.session_state.treatment=f"{name} Lung Cancer"
                st.session_state.doctor = doctors
                switch_page('Book_Appointment')
            
            st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
            first_run = st.session_state.get("first_run", True)

            if first_run:
                if st.button("Chat with AI Bot"):
                    st.session_state.first_run = False
                    runner()
            else:
                runner()

    else:
        st.success("Feel safe! No disease detected")
        st.markdown("##### Need more information? :speech_balloon:", unsafe_allow_html=False)
        first_run = st.session_state.get("first_run", True)

        if first_run:
            if st.button("Chat with AI Bot"):
                st.session_state.first_run = False
                runner()
        else:
            runner()
    
